import re
from collections import Counter, defaultdict
from docx import Document

def get_paragraph_info(para):
    run = para.runs[0] if para.runs else None
    return {
        "text": para.text.strip(),
        "font_size": run.font.size.pt if run and run.font.size else None,
        "bold": run.bold if run else False,
        "italic": run.italic if run else False,
        "color": str(run.font.color.rgb) if run and run.font.color and run.font.color.rgb else None,
        "alignment": para.alignment,
        "first_line_indent": para.paragraph_format.first_line_indent,
        "left_indent": para.paragraph_format.left_indent,
        "line_spacing": para.paragraph_format.line_spacing,
        "spacing_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
        "spacing_after": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
        "style_name": para.style.name,
        "font_name": run.font.name if run and run.font.name else None,
    }

def extract_references_from_doc(file_path):
    doc = Document(file_path)
    references = []
    found_references = False
    for para in doc.paragraphs:
        if para.text.strip().lower() == "references":
            found_references = True
            continue
        if found_references and para.text.strip():
            references.append(para.text.strip())
    return references

def extract_references(text):
    match = re.search(r"(References|REFERENCES)[\s\n]+(.+)", text, re.DOTALL)
    if match:
        refs_section = match.group(2)
        references = refs_section.strip().split("\n")
        return [r.strip() for r in references if len(r.strip()) > 30]
    return []

def check_duplicates(refs):
    return [item for item, count in Counter(refs).items() if count > 1]

def check_self_citations(refs, author_name=""):
    if not author_name:
        return []
    return [r for r in refs if author_name.lower() in r.lower()]

def check_qubahan(refs):
    return [r for r in refs if "Qubahan Academic Journal".lower() in r.lower()]

def check_apa_format(refs):
    bold_violations = []
    doi_violations = []
    for r in refs:
        if not re.search(r"\*\*\(\d{4}\)\*\*", r):
            bold_violations.append(r)
        if "doi.org" in r.lower():
            doi_violations.append(r)
    return bold_violations, doi_violations

def check_multiple_mentions(refs):
    author_counts = Counter()
    author_refs = defaultdict(list)

    for ref in refs:
        ref_clean = re.sub(r"^\[\d+\]\s*", "", ref)
        candidates = re.findall(r"\b[A-Z]\.\s*[A-Z][a-z]+", ref_clean)
        for name in candidates:
            norm_name = re.sub(r"\s+", " ", name).strip().lower()
            author_counts[norm_name] += 1
            author_refs[norm_name].append(ref)

    return {
        author: list(set(author_refs[author]))
        for author, count in author_counts.items()
        if count >= 4
    }

def check_missing_citations(text, refs):
    missing = []
    for i, ref in enumerate(refs):
        if f"[{i+1}]" not in text:
            missing.append(i + 1)
    return missing

def check_reference_formatting(file_path):
    doc = Document(file_path)
    issues = []
    found_references = False
    checking_refs = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if text == "References":
            found_references = True
            info = get_paragraph_info(para)

            if info["font_size"] != 10:
                issues.append("❌ 'References' heading should be 10 pt.")
            if info["color"] not in ["0000FF", "0000ff"]:
                issues.append("❌ 'References' heading should be blue (hex #0000FF).")
            if abs(info["spacing_before"] - 17) > 1:
                issues.append("❌ 'References' heading should have 17 pt spacing before.")
            if str(info["line_spacing"]).lower() != "single":
                issues.append("❌ 'References' heading should have single line spacing.")

            checking_refs = True
            continue

        if checking_refs:
            if not text:
                continue
            info = get_paragraph_info(para)

            if info["font_size"] != 8:
                issues.append(f"❌ Font size should be 8 pt in: '{text[:50]}'")
            if not info["font_name"] or "palatino" not in info["font_name"].lower():
                issues.append(f"❌ Font should be Palatino Linotype in: '{text[:50]}'")
            if info["alignment"] != 3:
                issues.append(f"❌ Reference not justified: '{text[:50]}'")
            if not info["left_indent"] or abs(info["left_indent"].inches - 0.25) > 0.05:
                issues.append(f"❌ Hanging indent should be 0.25 inch in: '{text[:50]}'")
            if info["spacing_before"] > 1 or info["spacing_after"] > 1:
                issues.append(f"❌ Spacing before/after text should be 0 in: '{text[:50]}'")
            if str(info["line_spacing"]).lower() != "single":
                issues.append(f"❌ Line spacing should be single in: '{text[:50]}'")

    if not found_references:
        issues.append("❌ 'References' heading not found in the document.")

    return issues

def check_references(text, author_name=""):
    results = {}
    refs = extract_references(text)
    if not refs:
        return {"error": "No references found. Ensure your document contains a 'References' section."}

    results["Total References"] = len(refs)
    results["Duplicate References"] = check_duplicates(refs)
    results["Self-Citations"] = check_self_citations(refs, author_name)
    results["Qubahan Citations"] = check_qubahan(refs)
    bold_violations, doi_violations = check_apa_format(refs)
    results["APA Style Violations"] = {
        "Missing Bold Year": bold_violations,
        "Contains DOI": doi_violations
    }
    results["Highly Cited Authors (≥4)"] = check_multiple_mentions(refs)
    results["Missing In-Text Citations"] = check_missing_citations(text, refs)
    results["Extracted References"] = refs
    return results
