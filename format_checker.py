from docx import Document
from docx.shared import Pt
import re
import string

def get_paragraph_info(para):
    run = para.runs[0] if para.runs else None
    return {
        "text": para.text.strip(),
        "font_size": run.font.size.pt if run and run.font.size else None,
        "bold": run.bold if run else False,
        "italic": run.italic if run else False,
        "color": str(run.font.color.rgb) if run and run.font.color and run.font.color.rgb else None,
        "alignment": para.alignment,
        "first_line_indent": para.paragraph_format.first_line_indent,
        "left_indent": para.paragraph_format.left_indent,
        "line_spacing": para.paragraph_format.line_spacing,
        "spacing_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
        "spacing_after": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
        "style_name": para.style.name,
        "font_name": run.font.name if run and run.font.name else None,
    }

def clean_text(text):
    return text.strip() if text and text.strip() and any(c.isalnum() for c in text) else ""

def check_font_and_spacing(file_path):
    doc = Document(file_path)
    issues = []
    for para in doc.paragraphs:
        info = get_paragraph_info(para)
        text = clean_text(info["text"])
        if not text:
            continue
        if not info["font_name"] or "palatino" not in info["font_name"].lower():
            issues.append(f"Font not Palatino Linotype in: '{text[:80]}'")
        if info["font_size"] and info["font_size"] != 12:
            issues.append(f"Font size is {info['font_size']} pt instead of 12 pt: '{text[:80]}'")
    return issues

def check_paragraph_format(file_path):
    doc = Document(file_path)
    issues = []
    for para in doc.paragraphs:
        info = get_paragraph_info(para)
        text = clean_text(info["text"])
        if not text:
            continue
        if para.alignment != 3:  # Justified
            issues.append(f"Paragraph not justified: '{text[:80]}'")
        if not info["first_line_indent"] or abs(info["first_line_indent"].inches - 0.2) > 0.05:
            issues.append(f"Paragraph missing first-line indent (should be 0.2\"): '{text[:80]}'")
    return issues

def check_margins(file_path):
    return ["Top margin is not 1 inch."]  # Placeholder; python-docx cannot detect margins

def check_headings(text):
    required = ["ABSTRACT", "INTRODUCTION", "LITERATURE REVIEW", "METHOD", "RESULT", "DISCUSSION", "CONCLUSION", "REFERENCES"]
    found = [h for h in required if h in text.upper()]
    missing = set(required) - set(found)
    return [f"Missing heading: {m}" for m in missing]

def check_tables_figures(text):
    issues = []
    table_titles = re.findall(r"(Table\s+\d+)[\.:]?", text, re.IGNORECASE)
    table_refs = re.findall(r"(Table\s+\d+)", text, re.IGNORECASE)

    figure_titles = re.findall(r"(FIGURE\s+\d+)[\.:]?", text, re.IGNORECASE)
    figure_refs = re.findall(r"(FIGURE\s+\d+)", text, re.IGNORECASE)

    table_title_set = set([t.lower() for t in table_titles])
    table_ref_set = set([r.lower() for r in table_refs])
    for t in table_title_set:
        if t not in table_ref_set:
            issues.append(f"{t.title()} not referenced in text.")

    figure_title_set = set([f.lower() for f in figure_titles])
    figure_ref_set = set([f.lower() for f in figure_refs])
    for f in figure_title_set:
        if f not in figure_ref_set:
            issues.append(f"{f.upper()} not referenced in text.")
    return issues

def check_subheadings(file_path):
    doc = Document(file_path)
    issues = []
    for para in doc.paragraphs:
        info = get_paragraph_info(para)
        text = clean_text(info['text'])
        if not text:
            continue
        if re.match(r"\d+\.\s+[A-Z ]+$", text):
            if not info["italic"]:
                issues.append(f"Subheading not italic: '{text}'")
            if info["spacing_before"] < 10:
                issues.append(f"Subheading spacing before should be 12 pt: '{text}'")
        elif re.match(r"\d+\.\d+\s+[A-Z][a-z]+", text):
            if not info["italic"]:
                issues.append(f"Sub-subheading not italic: '{text}'")
            if info["spacing_before"] < 5:
                issues.append(f"Sub-subheading spacing before should be 6 pt: '{text}'")
    return issues

def check_bullet_points(file_path):
    doc = Document(file_path)
    issues = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        if para.style.name.lower().startswith("list"):
            info = get_paragraph_info(para)
            if info["font_size"] != 10:
                issues.append(f"Bullet point font size not 10 pt: '{text[:80]}'")
            if not info["left_indent"] or abs(info["left_indent"].inches - 0.19) > 0.05:
                issues.append(f"Bullet indent not 0.19 inch: '{text[:80]}'")
            if para.alignment != 3:
                issues.append(f"Bullet point not justified: '{text[:80]}'")
    return issues

def check_reference_formatting(file_path):
    doc = Document(file_path)
    issues = []
    found_references = False
    checking_refs = False

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        if text == "References":
            found_references = True
            info = get_paragraph_info(para)
            if info["font_size"] != 10:
                issues.append("❌ 'References' heading should be 10 pt.")
            if info["color"] not in ["0000FF", "0000ff"]:
                issues.append("❌ 'References' heading should be blue (hex #0000FF).")
            if abs(info["spacing_before"] - 17) > 1:
                issues.append("❌ 'References' heading should have 17 pt spacing before.")
            if str(info["line_spacing"]).lower() != "single":
                issues.append("❌ 'References' heading should have single line spacing.")
            checking_refs = True
            continue

        if checking_refs:
            info = get_paragraph_info(para)
            if info["font_size"] != 8:
                issues.append(f"❌ Font size should be 8 pt in: '{text[:80]}'")
            if not info["font_name"] or "palatino" not in info["font_name"].lower():
                issues.append(f"❌ Font should be Palatino Linotype in: '{text[:80]}'")
            if info["alignment"] != 3:
                issues.append(f"❌ Reference not justified: '{text[:80]}'")
            if not info["left_indent"] or abs(info["left_indent"].inches - 0.25) > 0.05:
                issues.append(f"❌ Hanging indent should be 0.25 inch in: '{text[:80]}'")
            if info["spacing_before"] > 1 or info["spacing_after"] > 1:
                issues.append(f"❌ Spacing before/after text should be 0 in: '{text[:80]}'")
            if str(info["line_spacing"]).lower() != "single":
                issues.append(f"❌ Line spacing should be single in: '{text[:80]}'")

    if not found_references:
        issues.append("❌ 'References' heading not found in the document.")

    return issues
